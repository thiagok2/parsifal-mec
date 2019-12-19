# coding: utf-8

#from django.template import RequestContext
from django.conf import settings as djangoSettings

from docx import Document
from django.db.models import Count
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.utils.translation import ugettext as _
from parsifal.reviews.conducting.views import article_meta_analysis

from parsifal.reviews.models import Article, QualityAssessment, DataExtraction, DataExtractionField
from django.utils.html import escape
from parsifal.utils.svg_to_png import svg_to_png
from docx.shared import Inches

import xlsxwriter 

def export_review_to_docx(review, sections, request):
    document = Document()

    if 'name' in sections:
        h = document.add_heading(review.title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('')

    if 'authors' in sections:
        authors = list()
        authors.append(review.author.profile.get_screen_name())
        for author in review.co_authors.all():
            authors.append(author.profile.get_screen_name())
        p = document.add_paragraph(', '.join(authors))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('')

    if 'description' in sections:
        if review.description:
            document.add_paragraph(review.description)


    document.add_heading(_('Planning'), level=2)

    if review.objective:
        document.add_paragraph(review.objective)

    '''
        PICOC
    '''
    if 'picoc' in sections:

        if review.isPicoc():
            document.add_heading(_('PICOC'), level=3)
        elif review.isPicos():
            document.add_heading(_('PICOS'), level=3)
        elif review.isStudyTypeFree():
            document.add_heading(_('Free Text'), level=3)
        else:
            document.add_heading( review.study_type, level=3)

        if review.isStudyTypeFree():
            p = document.add_paragraph('', style='List Bullet')
            p.add_run(_('Study: ')).bold = True
            p.add_run(review.pico_text)
        else:
            p = document.add_paragraph('', style='List Bullet')
            p.add_run(_('Population: ')).bold = True
            p.add_run(review.population)

            p = document.add_paragraph('', style='List Bullet')
            p.add_run(_('Intervention: ')).bold = True
            p.add_run(review.intervention)

            p = document.add_paragraph('', style='List Bullet')
            p.add_run(_('Comparison: ')).bold = True
            p.add_run(review.comparison)

            p = document.add_paragraph('', style='List Bullet')
            p.add_run(_('Outcome: ')).bold = True
            p.add_run(review.outcome)

            if review.isPicoc():
                p = document.add_paragraph('', style='List Bullet')
                p.add_run(_('Context: ')).bold = True
                p.add_run(review.context)
            elif review.isPicos():
                p = document.add_paragraph('', style='List Bullet')
                p.add_run(_('Study Type: ')).bold = True
                p.add_run(review.study_type)
            else:
                p = document.add_paragraph('', style='List Bullet')
                p.add_run(_('Custom: ')).bold = True
                p.add_run(review.pico_text)


    '''
        Research Questions
    '''
    if 'research_questions' in sections:
        document.add_heading(_('Research Questions'), level=3)

        for question in review.research_questions.all():
            document.add_paragraph(question.question, style='List Number')

    '''
        Keywords and Synonym
    '''
    if 'keywords_synonyms' in sections:
        document.add_heading(_('Keywords and Synonyms'), level=3)

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = _('Keyword')
        hdr_cells[1].text = _('Synonyms')

        for keyword in review.get_keywords():
            row_cells = table.add_row().cells
            row_cells[0].text = keyword.description
            row_cells[1].text = ', '.join(keyword.synonyms.all().values_list('description', flat=True))

    '''
        Search String
    '''
    if 'search_string' in sections:
        document.add_heading(_('Search String'), level=3)
        document.add_paragraph(review.get_generic_search_string().search_string)

    '''
        Sources
    '''
    if 'sources' in sections:
        document.add_heading(_('Sources'), level=3)

        for source in review.sources.all():
            text = source.name
            if source.url:
                text = u'{0} ({1})'.format(source.name, source.url)
            document.add_paragraph(text, style='List Bullet')

    '''
        Risks
    '''
    if 'risks' in sections:
        document.add_heading(_('Risks To Review Validity'), level=3)

        for risk in review.get_risks():
            document.add_paragraph(risk.risk, style='List Bullet')

    '''
        Statistical Methods and Conventions
    '''
    if 'statistics_methods' in sections:
        document.add_heading(_('Statistical Methods and Conventions'), level=3)
        document.add_paragraph(review.statistical_methods)

    '''
        Selection Criteria
    '''
    if 'selection_criteria' in sections:
        document.add_heading(_('Selection Criteria'), level=3)

        p = document.add_paragraph()
        p.add_run(_('Inclusion Criteria:')).bold = True
        for criteria in review.get_inclusion_criterias():
            document.add_paragraph(criteria.description, style='List Bullet')

        p = document.add_paragraph()
        p.add_run(_('Exclusion Criteria:')).bold = True
        for criteria in review.get_exclusion_criterias():
            document.add_paragraph(criteria.description, style='List Bullet')

    '''
        Quality Assessment Checklist
    '''
    if 'quality_assessment_checklist' in sections:
        document.add_heading(_('Quality Assessment Checklist'), level=3)

        p = document.add_paragraph()
        p.add_run(_('Questions:')).bold = True
        for quality_question in review.get_quality_assessment_questions():
            document.add_paragraph(quality_question.description, style='List Bullet')
            for quality_answer in quality_question.get_answers():
                 document.add_paragraph(quality_answer.description+"("+str(quality_answer.weight)+")", style='ListBullet2')

    '''
        Data Extraction Form
    '''
    if 'data_extraction_form' in sections:
        document.add_heading(_('Data Extraction Form'), level=3)
        for field in review.get_data_extraction_fields():
            document.add_paragraph(field.description + '('+field.get_field_type_display()+')', style='List Bullet')

    '''
        Conducting
    '''

    document.add_heading(_('Conducting'), level=2)

    '''
        Digital Libraries Search Strings
    '''

    if 'source_search_strings' in sections:
        document.add_heading(_('Digital Libraries Search Strings'), level=3)
        for search_session in review.get_latest_source_search_strings():
            p = document.add_paragraph()
            p.add_run(u'{0}:'.format(search_session.source.name)).bold = True
            document.add_paragraph(search_session.search_string)
            document.add_paragraph()
        else:
            document.add_paragraph(_('Not specified'))

    if 'number_imported_studies' in sections:
        document.add_heading(_('Imported Studies'), level=3)
        for source in review.sources.all():
            p = document.add_paragraph(style='List Bullet')
            p.add_run(u'{0}: '.format(source.name)).bold = True
            count = review.article_set.filter(source=source).count()
            p.add_run(str(count))

    if 'number_study_selection_status' in sections:
        document.add_heading(_('Study Selection'), level=3)
        result_status = Article.objects.filter(review_id=review.id).values('status').order_by('status').annotate(count=Count('status'))

        for result in result_status:
            p = document.add_paragraph(style='List Bullet')

            label =  dict(Article.ARTICLE_STATUS).get(result['status'])

            p.add_run(u'{0}: '.format(label)).bold = True
            p.add_run(str(result['count']))

    if 'quality_assessment' in sections:
        document.add_heading(_('Quality Assessment'), level=3)

        document.add_heading(_('Question/Answer'), level=4)
        result_quality_assessment = QualityAssessment.objects.filter(article__review_id=review.id).values('question__description','answer__description').order_by().annotate(Count('question'), Count('answer') )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = _('Question')
        hdr_cells[1].text = _('Answer')
        hdr_cells[2].text = _('Count')

        for result in result_quality_assessment:
            row_cells = table.add_row().cells
            row_cells[0].text = result['question__description']
            row_cells[1].text = result['answer__description']
            row_cells[2].text = str(result['answer__count'])

        document.add_heading(_('Accepted Articles - Summary'), level=4)
        table_articles = document.add_table(rows=1, cols=2)
        hdr_cells_articles = table_articles.rows[0].cells
        hdr_cells_articles[0].text = _('Article')
        hdr_cells_articles[1].text = _('Score')
        for article in review.get_accepted_articles():
            row_cells = table_articles.add_row().cells
            row_cells[0].text = ''
            if article.title:
                row_cells[0].text += article.title
            if article.year:
                row_cells[0].text += '('+article.year+')'
            row_cells[1].text = str(article.get_score())

    if 'data_analysis' in sections:

        document.add_heading(_('Data Analysis'), level=3)
        try:
            if review.is_metaanalysis:
                conclusions = article_meta_analysis(review, request)
                svg_to_png(conclusions['forest_plot'], review)

                document.add_picture(djangoSettings.MEDIA_ROOT + '/forest_plot-'+ review.name +'.png', width=Inches(5.90))

                if(conclusions):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = _('Article')
                    hdr_cells[1].text = _('Effect Size')
                    hdr_cells[2].text = _('Effect')
                    for data in conclusions['conclusions']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = data['article']
                        row_cells[1].text = str(data['effect_size'])
                        row_cells[2].text = data['effect']
                else:
                    document.add_paragraph(_('Not specified'))
            else:
                document.add_paragraph(_('This review is not a meta analysis'))
        except Exception as e:
            print 'erro ', e
            document.add_paragraph(_('An exception occurred when generate data_analysis report'))
    return document

def export_review_to_xlsx(review, workbook):
    bold = workbook.add_format({'bold': True})
    bold.set_align('vjustify')
    
    cell_format = workbook.add_format()
    #cell_format.set_text_wrap()
    cell_format.set_align('vjustify')

    articles = review.get_finished_data_extraction_articles()


    # Quality Assessment Checklist
    sheet_checklist = workbook.add_worksheet(_('Quality Assessment Checklist'))
    sheet_checklist.write(0, 0,_('Q'), bold)
    sheet_checklist.write(0, 1,_('Question'), bold)
    sheet_checklist.write(0, 2,_('Answers'), bold)
    sheet_checklist.set_column(1, 1, 100)
    sheet_checklist.set_column(2, 2, 50)

    quality_questions = review.get_quality_assessment_questions()
    
    l = 1
    for question in quality_questions:
        sheet_checklist.write(l, 0, 'Q'+str(l))
        sheet_checklist.write(l, 1, question.description, cell_format)
        quality_answers = question.get_answers()
        values = []
        for answer in quality_answers:
            values.append(answer.description + '(' + str(answer.weight) + ')')
        sheet_checklist.write(l, 2, '; '.join(values), cell_format)
        l = l + 1

    # Quality Assessment Articles
    sheet_quality_assessment = workbook.add_worksheet(_('Quality Assessment'))
    sheet_quality_assessment.write(0, 0,_('Bibtex'), bold)
    sheet_quality_assessment.write(0, 1,_('Question'), bold)
    sheet_quality_assessment.set_column(0, 0, 20)
    sheet_quality_assessment.set_column(1, 1, 100)
    sheet_quality_assessment.set_column(2, 20, 20)

    for c in range(quality_questions.count()):
        sheet_quality_assessment.write(0, c+1, 'Q'+str(c), bold)

    l = 1
    for article in articles:
        sheet_quality_assessment.write(l, 0, article.bibtex_key)
        sheet_quality_assessment.write(l, 1, article.title, cell_format)

        quality_assessment = article.get_quality_assesment()
        q = 2
        for question in quality_questions:
            try:
                question_answer = quality_assessment.filter(question__id=question.id).get()
            except:
                question_answer = None
            if question_answer is not None:
                sheet_quality_assessment.write(l, q, question_answer.answer.description  + '(' + str(question_answer.answer.weight) + ')')
            q = q + 1

        l = l + 1
    
    # Data Extraction Form

    sheet_form = workbook.add_worksheet(_('Data Extraction Form'))
    sheet_form.write(0, 0,_('Description'), bold)
    sheet_form.write(0, 1,_('Type'), bold)
    sheet_form.write(0, 2,_('Values'), bold)
    sheet_form.set_column(0, 0, 50)
    sheet_form.set_column(1, 1, 30)
    sheet_form.set_column(2, 2, 100)

    fields = review.get_data_extraction_fields()
    l = 1
    for field in fields:
        sheet_form.write(l, 0, field.description)
        sheet_form.write(l, 1, field.get_field_type_display())
        if field.is_select_field():
            values = []
            for value in field.get_select_values():
                values.append(value.value)

            sheet_form.write(l, 2, '; '.join(values), cell_format)
        l = l + 1

    # Data Extraction: Articles

    sheet = workbook.add_worksheet(_('Data Extraction'))
    sheet.set_row(0, 30)
    sheet.set_column(0, 0, 20)
    sheet.set_column(1, 1, 90)
    sheet.set_column(2, 2, 40)
   
    sheet.set_column(3, 2 + fields.count(), 20)
    
    sheet.write(0, 0,_('Bibtex'), bold)
    sheet.write(0, 1,_('Article'), bold)
    c = 2
    for field in fields:
        sheet.write(0, c, field.description, cell_format)
        c = c + 1

    i = 1
    for article in articles:
        sheet.write(i, 0, article.bibtex_key)
        sheet.write(i, 1, article.title, cell_format)
        
        c = 2
        for field in fields:
            try:
                extraction = DataExtraction.objects.get(article=article, field=field)
                if field.field_type == DataExtractionField.DATE_FIELD:
                    sheet.write_datetime(i, c, extraction.get_value(), cell_format)
                if field.field_type == DataExtractionField.BOOLEAN_FIELD:
                    sheet.write_boolean(i, c, extraction.get_value(), cell_format)
                else:
                    sheet.write(i, c, extraction.get_value(), cell_format)
            except Exception, e:
                extraction = None
                #sheet.write(i, c, 'NA')
            c = c + 1
        
        i = i + 1  
    
    
    workbook.close()

    return workbook
